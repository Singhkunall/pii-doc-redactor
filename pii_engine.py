"""
pii_engine.py - Enhanced PII Detection & Format-Preserving DOCX Redaction Engine
Includes Regex Detectors, Heuristic Detectors, Deterministic Hash Mapper,
and Run-Level DOCX Formatting Preservation.
"""

import re
import hashlib
from io import BytesIO
from docx import Document

# ---------------------------------------------------------------------------
# 1. REGEX DETECTORS (structured PII)
# ---------------------------------------------------------------------------

EMAIL_RE = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')

# Covers Indian +91 formats AND generic US/international phone formats
PHONE_RE = re.compile(
    r'(?:\+91[\s\-]?\d{2,5}[\s\-]?\d{3,5}[\s\-]?\d{0,5})'
    r'|(?:\(?\b\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}\b)'
)

SSN_RE = re.compile(r'\b\d{3}-\d{2}-\d{4}\b')

CREDIT_CARD_RE = re.compile(
    r'\b(?:\d{4}[\s\-]){3}\d{4}\b|\b\d{13,16}\b'
)

IP_ADDRESS_RE = re.compile(
    r'\b(?:(?:25[0-5]|2[0-4]\d|1?\d{1,2})\.){3}(?:25[0-5]|2[0-4]\d|1?\d{1,2})\b'
)

# Date of Birth detector
DOB_RE = re.compile(
    r'(?:Date of [Bb]irth|DOB)\s*[:\-]?\s*'
    r'((?:\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})'
    r'|(?:[A-Z][a-z]+ \d{1,2},? \d{4}))'
)

REGEX_CATEGORIES = [
    ("DATE_OF_BIRTH", DOB_RE),   # must run before generic patterns
    ("EMAIL", EMAIL_RE),
    ("SSN", SSN_RE),
    ("CREDIT_CARD", CREDIT_CARD_RE),
    ("IP_ADDRESS", IP_ADDRESS_RE),
    ("PHONE", PHONE_RE),
]

# ---------------------------------------------------------------------------
# 2. HEURISTIC DETECTORS (unstructured PII)
# ---------------------------------------------------------------------------

COMPANY_SUFFIX_RE = re.compile(
    r'\b(?:[A-Z][A-Za-z&.\-]+(?:\s+(?:of|and|&|[A-Z][A-Za-z&.\-]+)){0,5}\s+'
    r'(?:Private Limited|Pvt\.?\s?Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Family Trust))\b'
)

_NAME_CHUNK = r"[A-Z][a-z']+(?:\s+[A-Z][a-z']+){1,3}"

NAME_TRIGGER_RE = re.compile(
    r'(?:Contact [Pp]erson[s]?|Company Secretary|Compliance Officer|'
    r'Mr\.|Mrs\.|Ms\.|Dr\.|Authorised [Ss]ignatory|'
    r'Chairman|Managing Director|Chief Executive Officer|Chief Financial Officer)'
    r'\s*[:\-]?\s*(' + _NAME_CHUNK + r')'
)

# Name appears BEFORE its role
NAME_BEFORE_ROLE_RE = re.compile(
    r'\b(' + _NAME_CHUNK + r')\s+is\s+(?:our|the)\s+(?:Company Secretary|'
    r'Compliance Officer|Managing Director|Chief Executive Officer|'
    r'Chief Financial Officer|Director|Chairman)\b'
)

# Names inside a comma/slash separated list
NAME_LIST_RE = re.compile(
    r'\b(' + _NAME_CHUNK + r')\b(?=\s*(?:,|/| and ))'
)

LEGAL_DEFINED_TERM_WORDS = {
    "Red", "Herring", "Prospectus", "Selling", "Shareholders", "Shareholder",
    "Regulations", "Regulation", "Managers", "Manager", "Managerial",
    "Personnel", "Insurance", "Companies", "Mutual", "Funds", "Fund",
    "Pension", "Anchor", "Investors", "Investor", "Price", "Band",
    "Statutory", "Auditors", "Auditor", "Promoter", "Promoters", "Group",
    "Book", "Running", "Lead", "Offer", "Board", "Committee", "Company",
    "India", "Limited", "Private", "Trust", "Bank", "Securities",
    "Exchange", "Exchanges", "Wealth", "Management", "Registrar",
    "Compliance", "Officer", "Officers", "Secretary", "Our", "The",
    "Members", "Market", "Codes", "Equity", "Shares", "Net", "Proceeds",
    "Key", "Senior", "Corporate", "Identity", "Number", "Financial",
    "Condition", "Risk", "Factors", "Business", "Government", "State",
    "Litigation", "Chartered", "Accountants", "Fresh", "Issue", "Registered",
    "Brokers", "Regulatory", "Proposed", "Capital", "Expenditure", "Supa",
    "Facility", "Application", "Master", "Circular", "Long", "Term", "Off",
    "Pallod", "Farms", "Depositories", "Act", "Bandra", "Kurla", "Complex",
    "Think", "Techno", "Campus", "Kanjurmarg", "Railway", "Station",
    "Escrow", "Accounts", "Pricing", "Date", "Prabhat", "Road", "Marg",
    "Deccan", "Gymkhana", "Rupees", "Rupee", "Technical", "Conventions",
    "United", "States", "Retail", "Portion", "Taluka", "Khed", "First",
    "Bidder", "East", "West", "North", "South", "Appasaheb", "Marathe",
    "Prabhadevi", "Backbay", "Reclamation", "Churchgate", "Village",
    "Birdewadi", "District", "Nagar", "Based", "Employee", "Benefits", 
    "Bid", "Lot", "Blocked", "Amount", "Buena", "Monte", "Building", "Process", 
    "Kamgar", "Sangathna", "Collecting", "Depository", "Participants", "Designated",
    "Intermediaries", "Earnings", "Before", "Faster", "Adoption",
    "Finalizing", "Independent", "Individual", "Bidders", "Trade",
    "Leasehold", "Obligations", "Sabha", "Discussion", "Mauje", "Palve",
    "Khurd", "Monitoring", "Agency", "Agreement", "Harshal", "Hall",
    "Parents", "Bunglow", "Pushpakamal", "Apartment", "Qualified",
    "Institutional", "Buyer", "Quality", "Control", "Sahara", "Hotel",
    "Sancheti", "Hospital", "Shivajinagar", "Peth", "Shubhkamal",
    "Leasing", "Signature", "Specified", "Locations", "Swedish",
    "Krona", "Underwriting", "Unified", "Payments", "Interface",
    "Working", "Interest", "European", "Union", "Gopalkrupa", "Kubera",
    "Executive", "Director", "Directors", "Freehold", "Land", "Ground",
    "Floor", "Chakan", "Industrial", "Area", "Internal", "Koregaon",
    "Park", "Onyx", "Tower", "Main", "Management's",
    "Society", "General", "Terms", "Model", "Colony", "Branch",
    "Refund", "Account", "Chambers", "Tara",
    "Material", "Contracts", "Documents", "Inspection",
    "Website", "Madhya", "Pradesh", "Manufacturing", "Unit", "Minal",
    "Residency", "Revision", "Form", "Sponsor", "Banks", "Tanishq",
    "Showroom", "Venture", "House", "Industry", "Data", "Provider",
}

def _is_probable_person_name(candidate):
    words = candidate.split()
    if not (2 <= len(words) <= 4):
        return False
    return not any(w in LEGAL_DEFINED_TERM_WORDS for w in words)

ADDRESS_TRIGGER_RE = re.compile(
    r'(?:Registered Office|Corporate Office|Address)\s*[:\-]\s*([^\n]{10,180}?'
    r'(?:\d{3}\s?\d{3}|\d{6})[^\n]{0,60})'
)

def _clean_candidate(s):
    return s.strip().strip(',/').strip()

def find_all_pii(text):
    """Return list of (start, end, category, matched_text) spans, sorted by start position."""
    if not text:
        return []
    spans = []

    for cat, pattern in REGEX_CATEGORIES:
        for m in pattern.finditer(text):
            if cat == "DATE_OF_BIRTH":
                spans.append((m.start(1), m.end(1), cat, m.group(1)))
            else:
                spans.append((m.start(), m.end(), cat, m.group(0)))

    for m in COMPANY_SUFFIX_RE.finditer(text):
        spans.append((m.start(), m.end(), "COMPANY_NAME", m.group(0)))

    for m in NAME_TRIGGER_RE.finditer(text):
        name = m.group(1)
        if name and _is_probable_person_name(name):
            s = m.start(1)
            spans.append((s, s + len(name), "PERSON_NAME", name))

    for m in NAME_BEFORE_ROLE_RE.finditer(text):
        name = m.group(1)
        if name and _is_probable_person_name(name):
            s = m.start(1)
            spans.append((s, s + len(name), "PERSON_NAME", name))

    for m in NAME_LIST_RE.finditer(text):
        name = _clean_candidate(m.group(1))
        if name and _is_probable_person_name(name):
            spans.append((m.start(1), m.start(1) + len(name), "PERSON_NAME", name))

    for m in ADDRESS_TRIGGER_RE.finditer(text):
        addr = m.group(1)
        s = m.start(1)
        spans.append((s, s + len(addr), "ADDRESS", addr))

    # Resolve overlaps: sort by start ascending, then length descending
    spans.sort(key=lambda x: (x[0], -(x[1] - x[0])))
    resolved = []
    last_end = -1
    for s, e, cat, val in spans:
        if s >= last_end:
            resolved.append((s, e, cat, val))
            last_end = e
    return resolved

# ---------------------------------------------------------------------------
# 3. DETERMINISTIC FAKE-VALUE GENERATOR
# ---------------------------------------------------------------------------

FIRST_NAMES = ["Arjun", "Meera", "Rahul", "Anika", "Vikram", "Priya", "Sanjay",
               "Divya", "Karan", "Neha", "Rohan", "Isha", "Aditya", "Kavya"]
LAST_NAMES = ["Sharma", "Verma", "Iyer", "Nair", "Kapoor", "Reddy", "Bose",
              "Chatterjee", "Malhotra", "Desai", "Joshi", "Rao"]
COMPANY_WORDS = ["Meridian", "Bluepeak", "Silverline", "Northgate", "Cobalt",
                 "Riverstone", "Orion", "Falconridge", "Everline", "Zenith"]
COMPANY_SUFFIXES = ["Private Limited", "Limited", "Industries Limited"]
STREET_NAMES = ["Lakeview Road", "Cedar Lane", "MG Road", "Sunrise Avenue",
                 "Harbour Street", "Maple Grove"]
CITIES = [("Nashik", "Maharashtra"), ("Indore", "Madhya Pradesh"),
          ("Coimbatore", "Tamil Nadu"), ("Vadodara", "Gujarat"),
          ("Kochi", "Kerala")]
EMAIL_DOMAINS = ["mailbox-example.com", "corp-example.net", "example-mail.com"]

def _seed(value, salt="default"):
    h = hashlib.md5((salt + "::" + value.lower().strip()).encode()).hexdigest()
    return int(h, 16)

class FakeMapper:
    def __init__(self, salt="default"):
        self.salt = salt
        self.mapping = {}  

    def get(self, category, original):
        cat_map = self.mapping.setdefault(category, {})
        key = original.strip()
        if key in cat_map:
            return cat_map[key]
        fake = self._generate(category, key)
        cat_map[key] = fake
        return fake

    def _generate(self, category, original):
        seed = _seed(original, self.salt + category)
        if category == "EMAIL":
            local = "user" + str(seed % 9000 + 1000)
            domain = EMAIL_DOMAINS[seed % len(EMAIL_DOMAINS)]
            return f"{local}@{domain}"
        if category == "PHONE":
            digits = [str((seed >> (4 * i)) % 10) for i in range(10)]
            return "+91 " + "".join(digits[:5]) + " " + "".join(digits[5:])
        if category == "SSN":
            return f"{(seed % 900) + 100}-{(seed >> 8) % 90 + 10}-{(seed >> 16) % 9000 + 1000}"
        if category == "CREDIT_CARD":
            groups = [str((seed >> (16 * i)) % 9000 + 1000) for i in range(4)]
            return " ".join(groups)
        if category == "IP_ADDRESS":
            return ".".join(str((seed >> (8 * i)) % 255) for i in range(4))
        if category == "DATE_OF_BIRTH":
            day = seed % 28 + 1
            month = (seed >> 6) % 12 + 1
            year = 1960 + (seed >> 10) % 45
            return f"{month:02d}/{day:02d}/{year}"
        if category == "PERSON_NAME":
            first = FIRST_NAMES[seed % len(FIRST_NAMES)]
            last = LAST_NAMES[(seed >> 8) % len(LAST_NAMES)]
            return f"{first} {last}"
        if category == "COMPANY_NAME":
            word = COMPANY_WORDS[seed % len(COMPANY_WORDS)]
            suffix = COMPANY_SUFFIXES[(seed >> 8) % len(COMPANY_SUFFIXES)]
            return f"{word} {suffix}"
        if category == "ADDRESS":
            street = STREET_NAMES[seed % len(STREET_NAMES)]
            city, state = CITIES[(seed >> 8) % len(CITIES)]
            pin = 400000 + (seed >> 16) % 99999
            return f"{street}, {city} - {pin}, {state}, India"
        return "[REDACTED]"

# ---------------------------------------------------------------------------
# 4. FORMAT-PRESERVING DOCX REDACTION LOGIC
# ---------------------------------------------------------------------------

def redact_text(text, mapper=None, category_filter=None):
    """Simple plain-text redactor returning redacted string."""
    if not text or not text.strip():
        return text
    if mapper is None:
        mapper = FakeMapper()
        
    spans = find_all_pii(text)
    if category_filter:
        spans = [s for s in spans if s[2] in category_filter]
        
    if not spans:
        return text
        
    redacted_text = text
    for s, e, cat, val in reversed(spans):
        fake_val = mapper.get(cat, val)
        redacted_text = redacted_text[:s] + fake_val + redacted_text[e:]
        
    return redacted_text

def redact_paragraph_preserve_runs(paragraph, mapper=None, category_filter=None):
    """
    Redacts PII in a python-docx Paragraph while preserving character-level formatting (runs).
    """
    full_text = paragraph.text
    if not full_text or not full_text.strip():
        return 0
        
    spans = find_all_pii(full_text)
    if category_filter:
        spans = [s for s in spans if s[2] in category_filter]
        
    if not spans or not paragraph.runs:
        return 0
        
    if mapper is None:
        mapper = FakeMapper()

    # Process spans from right to left (end of paragraph to start)
    spans.sort(key=lambda x: x[0], reverse=True)
    count = 0

    for s_idx, e_idx, cat, val in spans:
        fake_val = mapper.get(cat, val)
        count += 1
        
        # Build character index map to runs
        run_map = []  # list of (run_obj, char_offset_in_run) for each char in paragraph
        for run in paragraph.runs:
            for c_offset in range(len(run.text)):
                run_map.append((run, c_offset))
                
        if e_idx > len(run_map) or s_idx >= len(run_map):
            continue

        start_run, start_offset = run_map[s_idx]
        end_run, end_offset = run_map[e_idx - 1]

        if start_run == end_run:
            # Span is inside a single run
            orig_run_text = start_run.text
            start_run.text = orig_run_text[:start_offset] + fake_val + orig_run_text[end_offset + 1:]
        else:
            # Span crosses multiple runs
            # 1. Update start run with fake_val
            orig_start_text = start_run.text
            start_run.text = orig_start_text[:start_offset] + fake_val
            
            # 2. Clear intermediate runs
            in_between = False
            for r in paragraph.runs:
                if r == start_run:
                    in_between = True
                    continue
                if r == end_run:
                    break
                if in_between:
                    r.text = ""
                    
            # 3. Trim end run
            orig_end_text = end_run.text
            end_run.text = orig_end_text[end_offset + 1:]
            
    return count

def process_docx_preserve_formatting(file_source, output_path_or_stream, mapper=None, category_filter=None):
    """
    Reads a docx file (path or BytesIO), redacts paragraphs/tables preserving run styles,
    and saves to output (path or BytesIO).
    Returns dict of summary statistics.
    """
    doc = Document(file_source)
    if mapper is None:
        mapper = FakeMapper()

    total_redactions = 0

    def process_paragraphs(paragraphs):
        nonlocal total_redactions
        for p in paragraphs:
            if p.text.strip():
                total_redactions += redact_paragraph_preserve_runs(p, mapper, category_filter)

    # Body paragraphs
    process_paragraphs(doc.paragraphs)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # Headers & Footers
    for section in doc.sections:
        if section.header:
            process_paragraphs(section.header.paragraphs)
        if section.footer:
            process_paragraphs(section.footer.paragraphs)

    doc.save(output_path_or_stream)
    return {
        "total_redactions": total_redactions,
        "mapping": mapper.mapping
    }

if __name__ == "__main__":
    sample = "Contact Person: Arjun Sharma (Phone: +91 98765 43210, Email: arjun@example.com). Address: Registered Office: MG Road, Nashik - 422001, Maharashtra, India."
    print("Original Text:\n", sample)
    print("\nRedacted Text:\n", redact_text(sample))
