"""
test_engine.py - Verification script for pii_engine.py
Tests PII detection accuracy, format preservation in DOCX runs, and table cell redaction.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from pii_engine import find_all_pii, redact_text, process_docx_preserve_formatting, FakeMapper

def run_tests():
    print("==========================================")
    print("1. TESTING PLAIN TEXT DETECTION & REDACTION")
    print("==========================================")
    
    test_text = (
        "Red Herring Prospectus\n"
        "Contact Person: Rahul Kapoor (Managing Director)\n"
        "Email: rahul.kapoor@meridian-corp.com | Phone: +91 98200 12345 | SSN: 123-45-6789\n"
        "DOB: 12/05/1985\n"
        "Registered Office: Lakeview Road, Indore - 452001, Madhya Pradesh, India\n"
        "Company: Silverline Private Limited\n"
    )
    
    spans = find_all_pii(test_text)
    print(f"Found {len(spans)} PII entities:")
    for s, e, cat, val in spans:
        print(f" - [{cat}] '{val}' (at {s}:{e})")
        
    redacted = redact_text(test_text)
    print("\nRedacted Text Output:")
    print(redacted)
    
    assert "rahul.kapoor@meridian-corp.com" not in redacted, "Email not redacted!"
    assert "+91 98200 12345" not in redacted, "Phone not redacted!"
    assert "123-45-6789" not in redacted, "SSN not redacted!"
    assert "Rahul Kapoor" not in redacted, "Person name not redacted!"
    print("\nPlain text assertion tests PASSED!")

    print("\n==========================================")
    print("2. TESTING DOCX RUN FORMATTING PRESERVATION")
    print("==========================================")

    test_docx_path = "sample_test_doc.docx"
    output_docx_path = "sample_test_redacted.docx"

    # Create test document with specific run styling
    doc = Document()

    # Paragraph with mixed formatting runs
    p = doc.add_paragraph()
    r1 = p.add_run("Authorized Signatory: ")
    r1.font.size = Pt(12)
    
    r2 = p.add_run("Anika Sharma")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    r3 = p.add_run(" is our Compliance Officer. Contact at ")
    r4 = p.add_run("anika.sharma@falcon.com")
    r4.italic = True
    r4.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Table with formatting
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Role"
    hdr_cells[1].text = "Contact Details"

    row_cells = table.rows[1].cells
    row_cells[0].text = "Director"
    
    p_cell = row_cells[1].paragraphs[0]
    r_cell1 = p_cell.add_run("Vikram Verma")
    r_cell1.bold = True
    r_cell2 = p_cell.add_run(" - Phone: +91 91234 56789")
    r_cell2.font.size = Pt(10)

    doc.save(test_docx_path)
    print(f"Created test docx at '{test_docx_path}'")

    # Run format-preserving redaction
    result = process_docx_preserve_formatting(test_docx_path, output_docx_path)
    print(f"Redaction complete! Total redactions: {result['total_redactions']}")

    # Verify redacted docx content and run counts
    red_doc = Document(output_docx_path)
    p_red = red_doc.paragraphs[0]

    print("\nRedacted Paragraph Runs:")
    for idx, r in enumerate(p_red.runs):
        print(f" Run {idx}: '{r.text}' (Bold: {r.bold}, Italic: {r.italic}, Color: {r.font.color.rgb if r.font and r.font.color else None})")

    # Assertions
    assert "Anika Sharma" not in p_red.text, "Name 'Anika Sharma' still present!"
    assert "anika.sharma@falcon.com" not in p_red.text, "Email 'anika.sharma@falcon.com' still present!"

    # Verify run 1 is still bold!
    assert p_red.runs[1].bold is True, "Run 1 (Person Name) lost bold styling!"
    # Verify run 3 is still italic!
    assert p_red.runs[3].italic is True, "Run 3 (Email) lost italic styling!"

    print("\nDOCX Formatting Preservation Assertion tests PASSED!")
    print("==========================================")

    # Clean up test files
    if os.path.exists(test_docx_path):
        os.remove(test_docx_path)
    if os.path.exists(output_docx_path):
        os.remove(output_docx_path)

if __name__ == "__main__":
    run_tests()
