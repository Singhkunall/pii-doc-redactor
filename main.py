"""
main.py - FastAPI Backend Server for Format-Preserving DOCX PII Redactor
"""

import os
from io import BytesIO
from typing import List, Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from docx import Document

from pii_engine import (
    find_all_pii,
    redact_text,
    FakeMapper,
    process_docx_preserve_formatting
)

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="PII Detection & DOCX Redactor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

# Mount static directory for frontend
static_dir = os.path.join(os.path.dirname(__file__), "static")
if not os.path.exists(static_dir):
    os.makedirs(static_dir)

app.mount("/static", StaticFiles(directory=static_dir), name="static")

@app.get("/")
def read_root():
    index_path = os.path.join(static_dir, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return HTMLResponse("<h2>PII Redactor Backend API is running.</h2>")

def extract_docx_text(contents: bytes) -> str:
    """Extract all text from a docx file in order."""
    doc = Document(BytesIO(contents))
    text_blocks = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            text_blocks.append(p.text)
            
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_blocks.append(" | ".join(row_text))
                
    return "\n".join(text_blocks)

@app.post("/api/analyze")
async def analyze_file(
    file: UploadFile = File(...),
    salt: str = Form("default")
):
    """Analyze uploaded DOCX/TXT file for PII entities."""
    filename = file.filename or "uploaded_document"
    contents = await file.read()
    
    if filename.endswith(".docx"):
        try:
            full_text = extract_docx_text(contents)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read DOCX file: {str(e)}")
    else:
        try:
            full_text = contents.decode("utf-8", errors="ignore")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to decode text file: {str(e)}")
            
    spans = find_all_pii(full_text)
    mapper = FakeMapper(salt=salt)
    
    summary = {}
    details = []
    
    for idx, (s, e, cat, val) in enumerate(spans):
        summary[cat] = summary.get(cat, 0) + 1
        fake_val = mapper.get(cat, val)
        details.append({
            "id": idx + 1,
            "category": cat,
            "original_value": val,
            "fake_value": fake_val,
            "start": s,
            "end": e
        })
        
    return {
        "filename": filename,
        "total_pii_found": len(spans),
        "summary": summary,
        "details": details,
        "full_text": full_text
    }

@app.post("/api/preview")
async def preview_redaction(
    file: UploadFile = File(...),
    salt: str = Form("default"),
    categories: Optional[str] = Form(None)
):
    """Generates HTML snippet preview of original vs redacted document."""
    contents = await file.read()
    filename = file.filename or ""
    
    if filename.endswith(".docx"):
        full_text = extract_docx_text(contents)
    else:
        full_text = contents.decode("utf-8", errors="ignore")
        
    category_list = categories.split(",") if categories else None
    spans = find_all_pii(full_text)
    if category_list:
        spans = [s for s in spans if s[2] in category_list]
        
    mapper = FakeMapper(salt=salt)
    
    # Build Original HTML with highlighted PII
    orig_html = full_text
    # Replace from right to left
    spans_desc = sorted(spans, key=lambda x: x[0], reverse=True)
    
    redacted_html = full_text
    
    for s, e, cat, val in spans_desc:
        fake_val = mapper.get(cat, val)
        badge_orig = f'<mark class="pii-badge badge-{cat.lower()}" title="{cat}">{val}</mark>'
        badge_redacted = f'<mark class="pii-badge redacted-badge badge-{cat.lower()}" title="Redacted from: {val}">{fake_val}</mark>'
        
        orig_html = orig_html[:s] + badge_orig + orig_html[e:]
        redacted_html = redacted_html[:s] + badge_redacted + redacted_html[e:]
        
    return {
        "original_html": orig_html.replace("\n", "<br>"),
        "redacted_html": redacted_html.replace("\n", "<br>"),
        "total_count": len(spans)
    }

@app.post("/api/redact")
async def redact_file(
    file: UploadFile = File(...),
    salt: str = Form("default"),
    categories: Optional[str] = Form(None)
):
    """Processes uploaded DOCX or TXT file, applying PII redaction, and returns redacted file download."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")
        
    contents = await file.read()
    category_list = categories.split(",") if categories else None
    mapper = FakeMapper(salt=salt)
    out_filename = f"Redacted_{file.filename}"

    if file.filename.lower().endswith(".docx"):
        input_stream = BytesIO(contents)
        output_stream = BytesIO()
        try:
            stats = process_docx_preserve_formatting(
                input_stream,
                output_stream,
                mapper=mapper,
                category_filter=category_list
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX redaction processing failed: {str(e)}")
            
        output_stream.seek(0)
        return StreamingResponse(
            output_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{out_filename}"'}
        )
    else:
        # Plain text file redaction
        try:
            text_str = contents.decode("utf-8", errors="ignore")
            redacted_str = redact_text(text_str, mapper=mapper, category_filter=category_list)
            output_stream = BytesIO(redacted_str.encode("utf-8"))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Text redaction processing failed: {str(e)}")
            
        output_stream.seek(0)
        return StreamingResponse(
            output_stream,
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{out_filename}"'}
        )

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port)
